# ---------------------------------------------------------------------
# Program Name: ReadWriteGoogleSheets
#
# Purpose: This code reads from and writes to a Google Sheet
#
#
#
# Author: Thalia Edwards
#
# Date: 03/02/2026
#
# Edits by: Taha Chowdhury between 3-02-26 - 3-12-2026
# ------------------------------------------------------------------------
import gspread
from oauth2client.service_account import ServiceAccountCredentials
from docx import Document
from pathlib import Path
from datetime import datetime, timedelta
import hashlib
import json
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.enum.dml import MSO_THEME_COLOR_INDEX
import os


def get_resume_link(resume):
    """
    Extracts the URL from a Google Sheets hyperlink formula.
    
    Args:
        resume (str): A Google Sheets formula string containing a hyperlink.
        
    Returns:
        str: The extracted URL from the hyperlink.
    """
    return resume.split('"')[1]


def format_submission_date(raw_value):
    """
    Converts Google Sheets serial dates to MM/DD/YYYY format.
    Handles both numeric serial dates and pre-formatted date strings.
    
    Args:
        raw_value: A numeric Google Sheets serial date or date string.
        
    Returns:
        str: Formatted date string in MM/DD/YYYY format.
    """
    try:
        # Google Sheets uses serial date format (days since 12/30/1899)
        serial = float(raw_value)
        dt = datetime(1899, 12, 30) + timedelta(days=serial)
        return dt.strftime("%m/%d/%Y")
    except (ValueError, TypeError):
        # Fallback: attempt to parse if value is already a date string
        text = str(raw_value).strip()
        for fmt in ("%m/%d/%Y %H:%M:%S", "%m/%d/%Y"):
            try:
                return datetime.strptime(text, fmt).strftime("%m/%d/%Y")
            except ValueError:
                pass
        # If all parsing attempts fail, return original value
        return text

def add_hyperlink(paragraph, url, text):
    """
    Inserts a hyperlink into a Word document paragraph with proper formatting.
    Credit: ryan-rushton on GitHub for the original implementation.
    
    Args:
        paragraph: The docx paragraph object to insert the hyperlink into.
        url (str): The URL to link to.
        text (str): The visible text to display for the hyperlink.
        
    Returns:
        The run object containing the hyperlink.
    """
    # 1) Create relationship id between this paragraph and external URL
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    # 2) Build <w:hyperlink r:id="...">
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    hyperlink.set(qn("w:history"), "1")

    # 3) Build the visible run inside the hyperlink
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    # Apply Hyperlink style from Word template for consistency
    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)
    new_run.append(r_pr)

    # Add the visible text to display in the hyperlink
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)

    # 4) Attach hyperlink to the paragraph
    run = paragraph.add_run()
    run._r.append(hyperlink)

    # 5) Apply visual formatting as fallback if template lacks Hyperlink style
    run.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    run.font.underline = True

    return run


# =====================================================================
# CONFIGURATION AND AUTHENTICATION
# =====================================================================

# OAuth2 scopes required for Google Sheets and Drive API access
scopes = [
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive",
]

# Define project directory paths
SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_DIR = SCRIPT_DIR.parent  # pythonProject root directory
APPLICATION_DIR = PROJECT_DIR / "Applications"  # Base directory for generated application documents

# Path to service account credentials and document template
key_path = PROJECT_DIR / "secret_key" / "secret_key.json"
template_path = PROJECT_DIR / "Applications" / "Application Template.docx"

# Authenticate with Google using service account credentials
creds = ServiceAccountCredentials.from_json_keyfile_name(key_path, scopes=scopes)

# Connect to Google Sheets and retrieve application form responses
file = gspread.authorize(creds)
workbook = file.open("application-form-responses")
sheet = workbook.sheet1

# =====================================================================
# DATA RETRIEVAL FROM GOOGLE SHEETS
# =====================================================================

# Fetch all rows from the spreadsheet (preserves formulas for hyperlink extraction)
rows = sheet.get_all_values(value_render_option="FORMULA")

# =====================================================================
# APPLICANT REGISTRY AND DATA PROCESSING
# =====================================================================

# Load persistent applicant registry (maps email to unique ID) or create new one
registry_path = PROJECT_DIR / "applicants_registry.json"
try:
    with open(registry_path, "r") as f:
        applicants_registry = json.load(f)
except FileNotFoundError:
    # Initialize empty registry if file doesn't exist (e.g., first run)
    applicants_registry = {}  # {email: applicant_id}

# Nested dictionary to organize applications by applicant and position
# Structure: applicants[applicant_id][position_name] = application_data
applicants = {}

# Process each row from the Google Sheet (skip header row)
for row in rows[1:]:
    # Extract and parse application data from spreadsheet row
    application = {
        "submission": format_submission_date(row[0]),
        "name": row[1],
        "email": row[2],
        "position": row[3],
        "resume_URL": get_resume_link(str(row[4])),
        "employment_status": row[5],
        "prior": row[6],
        "hear_about": row[7],
        "reason_left": row[8],
        "current_employer": row[9],
        "availability": row[10],
        "phone": row[11],
        "preferred": row[12],
        "acknowledgement": row[13],
        "age": row[14],
    }
    
    # Normalize email for consistent lookup (lowercase and trimmed)
    email = row[2].lower().strip()

    # Assign persistent applicant ID or retrieve existing one
    if email not in applicants_registry:
        # Generate new applicant ID (format: APP_0001, APP_0002, etc.)
        new_id_num = len(applicants_registry) + 1
        applicant_id = f"APP_{new_id_num:04d}"
        applicants_registry[email] = applicant_id
    else:
        # Reuse existing ID for returning/duplicate applicants
        applicant_id = applicants_registry[email]

    # Extract application details for local processing
    position = application.get("position")
    applicant_name = application.get("name")

    # Create new applicant entry if this is the first application from this ID
    if applicant_id not in applicants:
        applicants[applicant_id] = {}

    # Store application data (overwrites if duplicate application for same position)
    applicants[applicant_id][position] = application

# Persist applicant registry to disk for next script execution
# This ensures consistent IDs across multiple runs
with open(registry_path, "w") as f:
    json.dump(applicants_registry, f, indent=2)


# =====================================================================
# DOCUMENT GENERATION AND FILE MANAGEMENT
# =====================================================================

# Iterate through all applicants and their position applications
for applicant_id, positions_dict in applicants.items():
    for position_name, application in positions_dict.items():
        applicant = application.get("name")
        
        # Load template document for this application
        document = Document(template_path)
        section = document.sections[0]
        footer = section.footer
        foot_para = footer.paragraphs[0]
        foot_para = foot_para.add_run(str(applicant))
        # --------- PLACEHOLDER REPLACEMENT ---------
        # Replace all placeholders with actual application data
        for key, value in application.items():
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        placeholder = f"[{key}]"
                        if placeholder in cell.text:
                            if key == "resume_URL":
                                # Special handling: Remove placeholder and insert clickable hyperlink
                                cell.text = cell.text.replace(placeholder, "")
                                add_hyperlink(cell.paragraphs[0], str(value), "View Resume")
                            else:
                                # Standard replacement for text fields
                                cell.text = cell.text.replace(placeholder, str(value))
            
                        
        
        # --------- PATH DETERMINATION AND DIRECTORY CREATION ---------
        # Determine output path based on position and create directory if needed
        # Map position names to their respective folders (known positions)
        if position_name == "Certified Nurse Assistant":
            output_path = APPLICATION_DIR / "Certified Nurse Assistant" / f"{applicant} - {position_name} Application.docx"
        elif position_name == "Patient Care Assistant":
            output_path = APPLICATION_DIR / "Patient Care Assistant" / f"{applicant} - {position_name} Application.docx"
        elif position_name == "Home Health Aide":
            output_path = APPLICATION_DIR / "Home Health Aide" / f"{applicant} - {position_name} Application.docx"
        elif position_name == "Licensed Practical Nurse":
            output_path = APPLICATION_DIR / "Licensed Practical Nurse" / f"{applicant} - {position_name} Application.docx"
        elif position_name == "Registered Nurse":
            output_path = APPLICATION_DIR / "Registered Nurse" / f"{applicant} - {position_name} Application.docx"
        elif position_name == "Accountant":
            output_path = APPLICATION_DIR / "Accountant" / f"{applicant} - {position_name} Application.docx"
        elif position_name == "Auditor":
            output_path = APPLICATION_DIR / "Auditor" / f"{applicant} - {position_name} Application.docx"
        elif position_name == "System Administrator":
            output_path = APPLICATION_DIR / "System Administrator" / f"{applicant} - {position_name} Application.docx"
        else:
            # Fallback: create folder for any unknown positions
            output_path = APPLICATION_DIR / position_name / f"{applicant} - {position_name} Application.docx"
        
        # --------- DOCUMENT SAVING WITH ERROR HANDLING ---------
        # Attempt to save document; create directory if it doesn't exist
        try:
            document.save(output_path)
            print(f"Saved document for {applicant} applying for {position_name}")
        except FileNotFoundError:
            # Directory doesn't exist; create it and retry save
            print(f"Directory for position '{position_name}' not found. Creating the directory and retrying...")
            mkdir_path = APPLICATION_DIR / position_name
            mkdir_path.mkdir(parents=True, exist_ok=True)
            output_path = mkdir_path / f"{applicant} - {position_name} Application.docx"
            document.save(output_path)
        except Exception as e:
            # Log any unexpected errors
            print(f"Error saving document for {applicant} - {position_name}: {e}")